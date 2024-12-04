import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches

def pdf_a_docx_con_formato(pdf_path, docx_path):
    """
    Convierte un PDF a DOCX intentando mantener el formato original.
    """
    doc = Document()  # Crea un archivo Word
    pdf_document = fitz.open(pdf_path)  # Abre el PDF

    for page_number, page in enumerate(pdf_document):
        # Agrega un encabezado para cada página
        doc.add_heading(f"Página {page_number + 1}", level=1)

        # Extrae bloques de texto
        blocks = page.get_text("blocks")
        for block in sorted(blocks, key=lambda b: (b[1], b[0])):  # Ordena por coordenadas
            texto = block[4].strip()
            if texto:
                doc.add_paragraph(texto)

        # Extrae imágenes
        for img_index, img in enumerate(page.get_images(full=True)):
            xref = img[0]
            base_image = pdf_document.extract_image(xref)
            image_bytes = base_image["image"]

            # Manejar la imagen en memoria
            image_stream = BytesIO(image_bytes)

            # Añadir la imagen al documento
            doc.add_paragraph(f"Imagen {img_index + 1} de la página {page_number + 1}:")
            doc.add_picture(image_stream, width=Inches(4))  # Ajusta el tamaño

    # Guarda el archivo Word
    doc.save(docx_path)
    print(f"Documento convertido guardado en: {docx_path}")

    
    
# Rutas de entrada y salida
ruta_pdf = r"C:\Users\luisd\Desktop\oposiciones\modulo2\632627042-Bloque-2-y-4-TAI.pdf"
ruta_docx = "documento.docx"

# Convertir PDF a DOCX
pdf_a_docx_con_formato(ruta_pdf, ruta_docx)
